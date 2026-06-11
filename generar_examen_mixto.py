import os
import json
import random
import requests

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK

# =========================================================
# TOKENS
# =========================================================

ORG_TOKEN = (os.environ.get("ORG_PAT") or "").strip()

BANK_TOKEN = (
    os.environ.get("BANK_PAT") or ""
).strip() or ORG_TOKEN

# =========================================================
# URLS BANCOS
# =========================================================

QUESTION_BANK_GRAFOS = (
    os.environ.get("QUESTION_BANK_GRAFOS") or ""
).strip()

QUESTION_BANK_VECTORIAL = (
    os.environ.get("QUESTION_BANK_VECTORIAL") or ""
).strip()

QUESTION_BANK_DOCUMENTAL = (
    os.environ.get("QUESTION_BANK_DOCUMENTAL") or ""
).strip()

# =========================================================
# CONFIG
# =========================================================

CONFIG_JSON = os.environ.get(
    "CONFIG_JSON_MIXTO",
    ""
)

ALUMNOS_CSV = os.environ.get(
    "ALUMNOS_CSV",
    ""
)

CONFIG = json.loads(CONFIG_JSON)

# =========================================================
# GLOBALES
# =========================================================

PREGUNTAS_USADAS = {}

# =========================================================
# LOAD JSON
# =========================================================

def cargar_banco(url, nombre):

    print(f"\n📦 Cargando banco: {nombre}")

    headers = {
        "Authorization": f"Bearer {BANK_TOKEN}",
        "Accept": "application/vnd.github.raw"
    }

    r = requests.get(url, headers=headers)

    print("📡 STATUS:", r.status_code)

    if r.status_code != 200:
        raise Exception(
            f"Error cargando banco {nombre}"
        )

    data = r.json()

    preguntas = data.get("ejercicios", [])

    for p in preguntas:
        p["banco"] = nombre

    print(
        f"✅ {nombre}: {len(preguntas)} preguntas"
    )

    return preguntas

# =========================================================
# CONTAR EXÁMENES
# =========================================================

def contar_examenes():

    contenido = (
        ALUMNOS_CSV
        .replace("\\n", "\n")
        .strip()
    )

    lines = [
        l.strip()
        for l in contenido.split("\n")
        if l.strip()
    ]

    total = max(len(lines) - 1, 1)

    print(f"👨‍🎓 Exámenes a generar: {total}")

    return total

# =========================================================
# SELECCIÓN
# =========================================================

def seleccionar_preguntas(
    banco_preguntas,
    config_banco
):

    seleccion = []

    for nivel, cantidad in (
        config_banco["cantidad"].items()
    ):

        subset = [
            p for p in banco_preguntas
            if p["nivel"].lower() == nivel.lower()
        ]

        if len(subset) < cantidad:

            raise Exception(
                f"No hay suficientes preguntas "
                f"{nivel} en {config_banco['nombre']}"
            )

        elegidas = random.sample(
            subset,
            cantidad
        )

        # ==========================================
        # PUNTAJE POR NIVEL
        # ==========================================

        for p in elegidas:

            p["puntaje"] = (
                config_banco["puntaje"][nivel]
            )

            p["origen"] = (
                config_banco["nombre"]
            )

        seleccion.extend(elegidas)

    return seleccion

# =========================================================
# RUBRICA RESUMIDA
# =========================================================

def rubrica_resumida(pregunta):

    criterios = (
        pregunta
        .get("respuesta", {})
        .get("criterios", [])
    )

    if not criterios:
        return "Sin rúbrica"

    return ", ".join(criterios[:3])

# =========================================================
# GENERAR DOCX EXÁMENES
# =========================================================

def generar_doc_examen(examen, examen_idx):

    doc = Document("template_examen.docx")

    doc.add_page_break()

    # =====================================================
    # ESTILO GENERAL
    # =====================================================

    style = doc.styles["Normal"]

    style.font.name = "Arial"
    style.font.size = Pt(11)

   
    # =====================================================
    # INICIO PREGUNTAS
    # =====================================================

    doc.add_paragraph("")

    # =====================================================
    # PREGUNTAS
    # =====================================================

    for i, pregunta in enumerate(examen, start=1):

        # =================================================
        # PREFIJO
        # =================================================

        if pregunta["origen"] == "documental":

            prefijo = "documental"

        elif pregunta["origen"] == "grafos":

            prefijo = "grafos"

        else:

            prefijo = pregunta["origen"]

        # =================================================
        # ENUNCIADO
        # =================================================

        texto = (
            f"{i}. "
            f"{prefijo}: "
            f"{pregunta['enunciado']} "
            f"({pregunta['puntaje']} punto"
            f"{'s' if pregunta['puntaje'] != 1 else ''}) "
            f"(id: {pregunta['id']})"
        )

        para = doc.add_paragraph()

        run = para.add_run(texto)

        run.bold = True
        run.font.size = Pt(11)

        # =================================================
        # QUERY BASE
        # =================================================

        query_base = pregunta.get(
            "query_base",
            ""
        )

        if query_base:

            q = doc.add_paragraph()

            qr = q.add_run(
                "\nQuery inicial:\n"
                f"{query_base}"
            )

            qr.font.name = "Courier New"
            qr.font.size = Pt(9)

        # =================================================
        # CÓDIGO BASE
        # =================================================

        codigo_base = pregunta.get(
            "codigo_base",
            ""
        )

        if codigo_base:

            c = doc.add_paragraph()

            cr = c.add_run(
                "\nCódigo base:\n"
                f"{codigo_base}"
            )

            cr.font.name = "Courier New"
            cr.font.size = Pt(9)

        # =================================================
        # DATASET
        # =================================================

        dataset = pregunta.get(
            "dataset",
            ""
        )

        if dataset:

            d = doc.add_paragraph()

            dr = d.add_run(
                "\nDataset:\n"
                f"{dataset}"
            )

            dr.font.name = "Courier New"
            dr.font.size = Pt(9)

        # =================================================
        # RÚBRICA
        # =================================================

        criterios = (
            pregunta
            .get("respuesta", {})
            .get("criterios", [])
        )

        query_respuesta = (
            pregunta
            .get("respuesta", {})
            .get("query")
        )

        if criterios:

            rubrica = (
                "Rúbrica: " +
                ", ".join(criterios[:3])
            )

        elif query_respuesta:

            rubrica = (
                "Rúbrica:\n"
                "• Lectura comprensiva y razonamiento correcto del enunciado (0.4 puntos) (  )\n"
                "• Uso correcto de sintaxis, operadores, funciones y estructuras (0.4 puntos)  (  )\n"
                "• Resultado correcto y coherente con lo solicitado (0.2 puntos) (  )"
            )

        else:

            rubrica = (
                "Rúbrica: correcta resolución"
            )

        r = doc.add_paragraph()

        rr = r.add_run(rubrica)

        rr.italic = True
        rr.font.size = Pt(9)

        # =================================================
        # ESPACIO RESPUESTA
        # =================================================

        nivel = pregunta["nivel"].lower()

        if nivel == "facil":

            espacio = 8

        elif nivel in [
            "medio",
            "intermedia"
        ]:

            espacio = 12

        else:

            espacio = 18

        for _ in range(espacio):

            doc.add_paragraph("")

    # =====================================================
    # GUARDAR
    # =====================================================

    filename = f"examen_{examen_idx + 1}.docx"

    doc.save(filename)

    print(f"✅ Generado: {filename}")

# =========================================================
# MAIN
# =========================================================

def main():

    print("🚀 Generador presencial mixto")

    banco_grafos = cargar_banco(
        QUESTION_BANK_GRAFOS,
        "grafos"
    )

    banco_vectoriales = cargar_banco(
        QUESTION_BANK_VECTORIAL,
        "vectoriales"
    )

    banco_documental = cargar_banco(
        QUESTION_BANK_DOCUMENTAL,
        "documental"
    )

    total_examenes = contar_examenes()

    examenes = []

    for _ in range(total_examenes):

        preguntas_finales = []

        for config_banco in CONFIG["bancos"]:

            nombre = config_banco["nombre"]

            if nombre == "grafos":

                banco_actual = banco_grafos

            elif nombre == "vectoriales":

                banco_actual = banco_vectoriales

            elif nombre == "documental":

                banco_actual = banco_documental

            else:

                raise Exception(
                    f"Banco desconocido: {nombre}"
                )

            seleccion = seleccionar_preguntas(
                banco_actual,
                config_banco
            )

            preguntas_finales.extend(
                seleccion
            )

        random.shuffle(preguntas_finales)

        for p in preguntas_finales:
            PREGUNTAS_USADAS[
                f"{p['banco']}-{p['id']}"
            ] = p

        examenes.append(
            preguntas_finales
        )

    for idx, examen in enumerate(examenes):
        generar_doc_examen(examen, idx)

    print("🎉 FIN")

# =========================================================
# START
# =========================================================

if __name__ == "__main__":
    main()