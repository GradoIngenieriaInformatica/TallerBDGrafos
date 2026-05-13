import os
import json
import random
import requests

from docx import Document
from docx.shared import Pt

# =========================================================
# CONFIG
# =========================================================

ORG_TOKEN = (os.environ.get("ORG_PAT") or "").strip()

BANK_TOKEN = (
    os.environ.get("BANK_PAT") or ""
).strip() or ORG_TOKEN

QUESTION_BANK_URL = (
    os.environ.get("QUESTION_BANK_VECTORIAL") or ""
).strip()

CONFIG_JSON = os.environ.get("CONFIG_JSON", "")

ALUMNOS_CSV = os.environ.get("ALUMNOS_CSV", "")

CONFIG = json.loads(CONFIG_JSON) if CONFIG_JSON else {}

RESULTADOS = []
PREGUNTAS_USADAS = {}

# =========================================================
# CARGA PREGUNTAS
# =========================================================

def cargar_preguntas():

    print("\n📦 Cargando preguntas...")

    headers = {
        "Authorization": f"Bearer {BANK_TOKEN}",
        "Accept": "application/vnd.github.raw"
    }

    r = requests.get(QUESTION_BANK_URL, headers=headers)

    print("📡 STATUS:", r.status_code)

    if r.status_code != 200:
        raise Exception(f"Error cargando preguntas: {r.text}")

    data = r.json()

    preguntas = data.get("ejercicios", [])

    print(f"✅ Preguntas cargadas: {len(preguntas)}")

    return preguntas

# =========================================================
# ALUMNOS
# =========================================================

def parsear_alumnos():

    contenido = ALUMNOS_CSV.replace("\\n", "\n").strip()

    lines = [l.strip() for l in contenido.split("\n") if l.strip()]

    delimiter = ";" if ";" in lines[0] else ","

    alumnos = []

    for line in lines[1:]:

        parts = [p.strip() for p in line.split(delimiter)]

        # columna 0 = nombre
        if len(parts) >= 1:

            nombre = parts[0]

            alumnos.append(nombre)

    print(f"👨‍🎓 TOTAL alumnos: {len(alumnos)}")

    return alumnos

# =========================================================
# SELECCIÓN FIJA 8 PREGUNTAS
# =========================================================

def seleccionar_preguntas(preguntas):

    faciles = [p for p in preguntas if p["nivel"] == "facil"]
    medias = [
        p for p in preguntas
        if p["nivel"].lower() in ["medio", "media", "intermedia"]
    ]
    dificiles = [p for p in preguntas if p["nivel"] == "dificil"]

    if len(faciles) < 3 or len(medias) < 3 or len(dificiles) < 2:
        raise Exception("No hay suficientes preguntas por nivel")

    seleccion = (
        random.sample(faciles, 3) +
        random.sample(medias, 3) +
        random.sample(dificiles, 2)
    )

    random.shuffle(seleccion)

    return seleccion

# =========================================================
# FORMATO ENUNCIADO
# =========================================================

def formatear_enunciado(p, index):

    # asignación de puntos por nivel
    puntos = {
        "facil": 1,
        "intermedia": 1,
        "dificil": 2
    }.get(p["nivel"], 1)

    return (
        f"{index}. {p['enunciado']} "
        f"(ID: {p['id']}) "
        f"[{puntos} punto{'s' if puntos > 1 else ''}]"
    )
    
# =========================================================
# GENERAR DOC ÚNICO (CLAVE)
# =========================================================

def generar_doc_unico(asignaciones):

    doc = Document()

    for idx, (alumno, preguntas) in enumerate(asignaciones):

        # ==========================
        # TÍTULO
        # ==========================

        doc.add_heading(
            "EXAMEN — BASES DE DATOS VECTORIALES",
            level=1
        )

        # ==========================
        # ALUMNO + FIRMA
        # ==========================

        p_nombre = doc.add_paragraph()

        run = p_nombre.add_run(f"Alumno: {alumno}\n")
        run.bold = True
        run.font.size = Pt(12)

        firma = p_nombre.add_run(
            "\nFirma: _______________________________\n"
        )

        firma.font.size = Pt(11)

        doc.add_paragraph("")

        # ==========================
        # PREGUNTAS
        # ==========================

        for i, p in enumerate(preguntas):

            texto = formatear_enunciado(p, i + 1)

            para = doc.add_paragraph()

            run = para.add_run(texto)
            run.font.size = Pt(11)

            doc.add_paragraph(
                "\nRespuesta:\n"
                "\n"
                "\n"
                "\n"
                "\n"
                "\n"
            )

        # =====================================
        # PÁGINA BLANCA PARA IMPRESIÓN DÚPLEX
        # =====================================

        doc.add_page_break()
        doc.add_page_break()

    doc.save("examenes_vectoriales.docx")

    print("✅ DOC FINAL GENERADO: examenes_vectoriales.docx")
    
# =========================================================
# MAIN
# =========================================================

def main():

    print("🚀 Generador exámenes vectoriales")

    preguntas = cargar_preguntas()

    alumnos = parsear_alumnos()

    asignaciones = []

    for alumno in alumnos:

        seleccion = seleccionar_preguntas(preguntas)

        ids = [p["id"] for p in seleccion]

        RESULTADOS.append({
            "alumno": alumno,
            "examen": ids
        })

        for p in seleccion:
            PREGUNTAS_USADAS[p["id"]] = p

        asignaciones.append((alumno, seleccion))

        print(f"📦 {alumno} → {ids}")

    generar_doc_unico(asignaciones)

    # JSON resultado
    with open("examenes_vectoriales_generados.json", "w", encoding="utf-8") as f:
        json.dump(RESULTADOS, f, indent=2, ensure_ascii=False)

    print("🎉 FIN")

# =========================================================
# START
# =========================================================

if __name__ == "__main__":
    main()